"""
导出接口 - 支持报告和图表导出
支持PDF、Word、Markdown格式导出
"""

from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
import os
import tempfile
from typing import Optional, Literal
from fastapi import APIRouter, HTTPException, Depends, Query, Path
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from sqlalchemy.orm import selectinload
import os
import tempfile

from app.db.database import get_db
from app.db.models import ResearchTask, Chart
from app.core.logging import logger

router = APIRouter()


@router.get("/tasks/{task_id}/export/{format}")
async def export_report(
    task_id: int,
    format: Literal["pdf", "word", "markdown"] = Path(..., description="导出格式"),
    include_charts: bool = Query(True, description="是否包含图表"),
    db: AsyncSession = Depends(get_db),
):
    """
    导出研究报告
    支持PDF、Word、Markdown格式，包含图表
    """
    # 获取任务详情
    query = (
        select(ResearchTask)
        .options(
            selectinload(ResearchTask.plan_items),
            selectinload(ResearchTask.sources),
            selectinload(ResearchTask.agent_logs),
            selectinload(ResearchTask.charts),
        )
        .filter(ResearchTask.id == task_id)
    )
    result = await db.execute(query)
    task = result.scalar_one_or_none()

    if not task:
        raise HTTPException(status_code=404, detail="研究任务不存在")

    if not task.report_content and not task.summary:
        raise HTTPException(status_code=400, detail="报告内容为空，无法导出")

    try:
        if format == "markdown":
            return await _export_markdown(task, include_charts)
        elif format == "pdf":
            return await _export_pdf(task, include_charts)
        elif format == "word":
            return await _export_word(task, include_charts)
        else:
            raise HTTPException(status_code=400, detail="不支持的导出格式")

    except Exception as e:
        logger.error(f"导出失败: {e}")
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


async def _export_markdown(task: ResearchTask, include_charts: bool) -> FileResponse:
    """导出Markdown格式"""
    content = []

    # 标题
    content.append(f"# {task.query}")
    content.append("")

    # 元信息
    content.append("---")
    content.append(f"生成时间: {task.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    if task.completed_at:
        content.append(f"完成时间: {task.completed_at.strftime('%Y-%m-%d %H:%M:%S')}")
    content.append("---")
    content.append("")

    # 执行摘要
    if task.summary:
        content.append("## 执行摘要")
        content.append("")
        content.append(task.summary)
        content.append("")

    # 图表部分
    if include_charts and task.charts:
        content.append("## 数据图表")
        content.append("")
        for chart in sorted(task.charts, key=lambda x: x.order):
            content.append(f"### {chart.title}")
            content.append("")
            if chart.description:
                content.append(f"*{chart.description}*")
                content.append("")

            # 图表数据描述
            content.append("**图表数据：**")
            content.append("")
            content.append(f"- 图表类型: {chart.chart_type}")
            content.append(f"- 数据来源: {chart.created_by_agent}")

            # 简单的数据展示
            if chart.chart_type == "bar" and "categories" in chart.data:
                content.append("- 数据统计:")
                for i, category in enumerate(chart.data["categories"]):
                    value = (
                        chart.data.get("series", [{}])[0].get("data", [])[i]
                        if chart.data.get("series")
                        else 0
                    )
                    content.append(f"  - {category}: {value}")

            content.append("")

    # 主要内容
    if task.report_content:
        content.append(task.report_content)
        content.append("")

    # 参考来源
    if task.sources:
        content.append("## 参考来源")
        content.append("")
        for i, source in enumerate(task.sources, 1):
            content.append(f"{i}. **{source.title}**")
            if source.url:
                content.append(f"   - 链接: {source.url}")
            if source.content:
                content.append(f"   - 摘要: {source.content[:200]}...")
            content.append("")

    # 创建临时文件
    with tempfile.NamedTemporaryFile(
        mode="w", encoding="utf-8", suffix=".md", delete=False
    ) as f:
        f.write("\n".join(content))
        temp_file = f.name

    filename = f"research_report_{task.id}.md"
    return FileResponse(
        temp_file,
        media_type="text/markdown",
        filename=filename,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


async def _export_pdf(task: ResearchTask, include_charts: bool) -> FileResponse:
    """导出PDF格式（支持中文）"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Image,
        )
        from reportlab.lib.utils import ImageReader
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from io import BytesIO
        import matplotlib.pyplot as plt
        import matplotlib

        matplotlib.use("Agg")  # 使用非GUI后端
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="PDF导出功能需要安装reportlab和matplotlib，请联系管理员",
        )

    # ========== 注册中文字体 ==========
    # 尝试多个常见的中文字体路径（Windows/Linux/Mac）
    chinese_font_paths = [
        # Windows 字体路径
        "C:/Windows/Fonts/msyh.ttc",  # 微软雅黑
        "C:/Windows/Fonts/simsun.ttc",  # 宋体
        "C:/Windows/Fonts/simhei.ttf",  # 黑体
        "C:/Windows/Fonts/STFANGSO.TTF",  # 华文仿宋
        # Linux 字体路径
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        # Mac 字体路径
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]

    chinese_font_name = None
    for font_path in chinese_font_paths:
        if os.path.exists(font_path):
            try:
                font_name = "ChineseFont"
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                chinese_font_name = font_name
                logger.info(f"成功注册中文字体: {font_path}")
                break
            except Exception as e:
                logger.warning(f"注册字体失败 {font_path}: {e}")
                continue

    if not chinese_font_name:
        logger.warning("未找到可用的中文字体，PDF可能显示乱码")
        chinese_font_name = "Helvetica"  # 回退到默认字体

    # 创建临时PDF文件
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        temp_file = f.name

    # 创建PDF文档
    doc = SimpleDocTemplate(
        temp_file,
        pagesize=A4,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50,
    )

    # ========== 创建支持中文的自定义样式 ==========
    title_style = ParagraphStyle(
        "ChineseTitle",
        fontName=chinese_font_name,
        fontSize=18,
        leading=24,
        spaceAfter=20,
        alignment=TA_CENTER,
    )

    heading_style = ParagraphStyle(
        "ChineseHeading",
        fontName=chinese_font_name,
        fontSize=14,
        leading=20,
        spaceAfter=12,
        spaceBefore=16,
    )

    heading3_style = ParagraphStyle(
        "ChineseHeading3",
        fontName=chinese_font_name,
        fontSize=12,
        leading=18,
        spaceAfter=8,
        spaceBefore=12,
    )

    normal_style = ParagraphStyle(
        "ChineseNormal",
        fontName=chinese_font_name,
        fontSize=10,
        leading=16,
        spaceAfter=6,
    )

    meta_style = ParagraphStyle(
        "ChineseMeta",
        fontName=chinese_font_name,
        fontSize=9,
        leading=14,
        textColor="grey",
        alignment=TA_CENTER,
    )

    story = []

    # 直接展示报告内容（不显示标题和元信息，报告本身已包含）
    if task.report_content:
        # 分段处理长文本，避免单个Paragraph过长
        paragraphs = _split_report_to_paragraphs(task.report_content)
        for para_text, para_type in paragraphs:
            if para_type == "h1":
                story.append(Paragraph(_escape_xml(para_text), title_style))
            elif para_type == "h2":
                story.append(Paragraph(_escape_xml(para_text), heading_style))
            elif para_type == "h3":
                story.append(Paragraph(_escape_xml(para_text), heading3_style))
            elif para_type == "normal":
                story.append(
                    Paragraph(_escape_xml_preserve_br(para_text), normal_style)
                )
        story.append(Spacer(1, 16))

    # 图表部分
    if include_charts and task.charts:
        story.append(Paragraph("数据图表", heading_style))
        story.append(Spacer(1, 8))

        for chart in sorted(task.charts, key=lambda x: x.order):
            story.append(Paragraph(_escape_xml(chart.title), heading3_style))

            if chart.description:
                story.append(
                    Paragraph(f"<i>{_escape_xml(chart.description)}</i>", normal_style)
                )
                story.append(Spacer(1, 4))

            # 生成图表图片并插入
            try:
                chart_img_data = await _generate_chart_image(chart)
                if chart_img_data:
                    img_reader = ImageReader(BytesIO(chart_img_data))
                    img = Image(img_reader)
                    img.drawHeight = 250
                    img.drawWidth = 450
                    story.append(img)
                    story.append(Spacer(1, 12))
                else:
                    story.append(Paragraph("[图表生成失败]", normal_style))
            except Exception as e:
                logger.warning(f"生成图表图片失败: {e}")
                story.append(Paragraph("[图表生成失败]", normal_style))

    # 参考来源
    if task.sources:
        story.append(Paragraph("参考来源", heading_style))
        story.append(Spacer(1, 8))

        for i, source in enumerate(task.sources, 1):
            source_text = f"{i}. {_escape_xml(source.title)}"
            if source.url:
                source_text += f"<br/>链接: {_escape_xml(source.url)}"
            if source.content:
                content_preview = (
                    source.content[:150] + "..."
                    if len(source.content) > 150
                    else source.content
                )
                source_text += f"<br/>摘要: {_escape_xml(content_preview)}"
            story.append(Paragraph(source_text, normal_style))
            story.append(Spacer(1, 4))

    # 生成PDF
    doc.build(story)

    filename = f"research_report_{task.id}.pdf"
    return FileResponse(
        temp_file,
        media_type="application/pdf",
        filename=filename,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


def _escape_xml(text: str) -> str:
    """转义XML特殊字符，避免ReportLab解析错误"""
    if not text:
        return ""
    # 替换XML特殊字符
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    text = text.replace("'", "&apos;")
    return text


def _escape_xml_preserve_br(text: str) -> str:
    """转义XML特殊字符，但保留<br/>标签"""
    if not text:
        return ""
    # 先把<br/>临时替换掉
    text = text.replace("<br/>", "[[BR]]")
    text = text.replace("<br>", "[[BR]]")
    # 转义其他特殊字符
    text = _escape_xml(text)
    # 恢复<br/>
    text = text.replace("[[BR]]", "<br/>")
    return text


def _split_report_to_paragraphs(content: str) -> list:
    """将报告内容拆分为段落列表，识别标题和正文"""
    import re

    result = []

    # 按换行分割
    lines = content.split("\n")
    current_para = []

    for line in lines:
        stripped = line.strip()

        # 识别Markdown标题（按照标题级别区分）
        if stripped.startswith("# ") and not stripped.startswith("## "):
            # 先保存之前的段落
            if current_para:
                result.append(("\n".join(current_para), "normal"))
                current_para = []
            result.append((stripped[2:], "h1"))
        elif stripped.startswith("## ") and not stripped.startswith("### "):
            if current_para:
                result.append(("\n".join(current_para), "normal"))
                current_para = []
            result.append((stripped[3:], "h2"))
        elif stripped.startswith("### "):
            if current_para:
                result.append(("\n".join(current_para), "normal"))
                current_para = []
            result.append((stripped[4:], "h3"))
        elif stripped == "":
            # 空行分隔段落
            if current_para:
                result.append(("\n".join(current_para), "normal"))
                current_para = []
        else:
            # 普通文本，去掉Markdown格式符号
            clean_line = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)  # 去掉粗体
            clean_line = re.sub(r"\*(.*?)\*", r"\1", clean_line)  # 去掉斜体
            clean_line = re.sub(r"^\- ", "• ", clean_line)  # 列表符号
            clean_line = re.sub(r"^\d+\. ", "", clean_line)  # 有序列表
            current_para.append(clean_line)

    # 保存最后一个段落
    if current_para:
        result.append(("\n".join(current_para), "normal"))

    return result


async def _export_word(task: ResearchTask, include_charts: bool) -> FileResponse:
    """导出Word格式"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import matplotlib.pyplot as plt
        import io
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="Word导出功能需要安装python-docx和matplotlib，请联系管理员",
        )

    # 创建Word文档
    doc = Document()

    # 标题
    title = doc.add_heading(task.query, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 元信息
    doc.add_paragraph(f"生成时间: {task.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
    if task.completed_at:
        doc.add_paragraph(
            f"完成时间: {task.completed_at.strftime('%Y-%m-%d %H:%M:%S')}"
        )

    # 执行摘要
    if task.summary:
        doc.add_heading("执行摘要", 1)
        doc.add_paragraph(task.summary)

    # 图表部分
    if include_charts and task.charts:
        doc.add_heading("数据图表", 1)

        for chart in sorted(task.charts, key=lambda x: x.order):
            doc.add_heading(chart.title, 2)
            if chart.description:
                doc.add_paragraph(chart.description, style="Intense Quote")

            # 生成图表图片并插入
            try:
                chart_img_data = await _generate_chart_image(chart)
                if chart_img_data:
                    # 保存图片到临时文件
                    with tempfile.NamedTemporaryFile(
                        suffix=".png", delete=False
                    ) as img_file:
                        img_file.write(chart_img_data)
                        img_temp_file = img_file.name

                    doc.add_picture(img_temp_file, width=Inches(6))
                    os.unlink(img_temp_file)  # 删除临时文件
                else:
                    doc.add_paragraph("[图表生成失败]")
            except Exception as e:
                logger.warning(f"生成图表图片失败: {e}")
                doc.add_paragraph("[图表生成失败]")

    # 主要内容
    if task.report_content:
        doc.add_heading("报告内容", 1)
        doc.add_paragraph(task.report_content)

    # 参考来源
    if task.sources:
        doc.add_heading("参考来源", 1)
        for i, source in enumerate(task.sources, 1):
            p = doc.add_paragraph(f"{i}. {source.title}")
            if source.url:
                p.add_run(f"\n链接: {source.url}")
            if source.content:
                p.add_run(f"\n摘要: {source.content[:200]}...")

    # 保存到临时文件
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        temp_file = f.name
        doc.save(temp_file)

    filename = f"research_report_{task.id}.docx"
    return FileResponse(
        temp_file,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


async def _generate_chart_image(chart: Chart) -> Optional[bytes]:
    """生成图表图片"""
    try:
        import matplotlib.pyplot as plt
        import io
        import numpy as np

        plt.figure(figsize=(8, 6))
        data = chart.data

        if chart.chart_type == "bar":
            if "categories" in data and "series" in data:
                categories = data["categories"]
                series = data["series"]
                if isinstance(series, list) and len(series) > 0:
                    series_data = series[0].get("data", [])
                    plt.bar(categories, series_data, color="#3b82f6")
                    plt.title(chart.title, fontsize=14, pad=20)
                    plt.xticks(rotation=45, ha="right")
                    plt.ylabel("数量")

        elif chart.chart_type == "line":
            if "categories" in data and "series" in data:
                categories = data["categories"]
                series = data["series"]
                if isinstance(series, list) and len(series) > 0:
                    series_data = series[0].get("data", [])
                    plt.plot(
                        categories,
                        series_data,
                        marker="o",
                        linewidth=2,
                        color="#3b82f6",
                    )
                    plt.title(chart.title, fontsize=14, pad=20)
                    plt.xticks(rotation=45, ha="right")
                    plt.grid(True, alpha=0.3)

        elif chart.chart_type == "pie":
            labels = data.get("labels", [])
            sizes = data.get("series", [])
            if labels and sizes and len(labels) == len(sizes):
                colors = [
                    "#3b82f6",
                    "#10b981",
                    "#f59e0b",
                    "#ef4444",
                    "#8b5cf6",
                    "#06b6d4",
                ]
                plt.pie(
                    sizes,
                    labels=labels,
                    colors=colors[: len(labels)],
                    autopct="%1.1f%%",
                    startangle=90,
                )
                plt.title(chart.title, fontsize=14, pad=20)
                plt.axis("equal")

        elif chart.chart_type == "radar":
            if "series" in data and data["series"]:
                series = data["series"]
                if isinstance(series, list) and len(series) > 0:
                    values = series[0].get("data", [])
                    if values:
                        # 为雷达图创建角度
                        num_vars = len(values)
                        angles = [
                            n / float(num_vars) * 2 * np.pi for n in range(num_vars)
                        ]
                        angles += angles[:1]  # 闭合图形
                        values += values[:1]

                        fig, ax = plt.subplots(
                            figsize=(6, 6), subplot_kw=dict(projection="polar")
                        )
                        ax.plot(angles, values, "o-", linewidth=2, color="#3b82f6")
                        ax.fill(angles, values, alpha=0.25, color="#3b82f6")
                        ax.set_title(chart.title, fontsize=14, pad=20)

                        # 设置标签
                        if "indicators" in data:
                            labels = [
                                ind.get("name", f"指标{i}")
                                for i, ind in enumerate(data["indicators"][:num_vars])
                            ]
                            ax.set_xticks(angles[:-1])
                            ax.set_xticklabels(labels)

                        plt.close("all")  # 关闭所有图形
                        buf = io.BytesIO()
                        fig.savefig(buf, format="png", bbox_inches="tight", dpi=150)
                        buf.seek(0)
                        return buf.getvalue()

        # 对于其他图表类型，保存图片到内存
        if chart.chart_type in ["bar", "line", "pie"]:
            plt.tight_layout()
            buf = io.BytesIO()
            plt.savefig(buf, format="png", bbox_inches="tight", dpi=150)
            plt.close("all")
            buf.seek(0)
            return buf.getvalue()

        return None

    except Exception as e:
        logger.error(f"生成图表图片失败: {e}")
        import traceback

        logger.error(traceback.format_exc())
        return None


def _markdown_to_simple_html(markdown_text: str) -> str:
    """简单的Markdown到HTML转换"""
    import re

    # 标题转换
    html = re.sub(r"^### (.*)$", r"<h3>\1</h3>", markdown_text, flags=re.MULTILINE)
    html = re.sub(r"^## (.*)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"^# (.*)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)

    # 粗体
    html = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", html)
    html = re.sub(r"__(.*?)__", r"<b>\1</b>", html)

    # 斜体
    html = re.sub(r"\*(.*?)\*", r"<i>\1</i>", html)
    html = re.sub(r"_(.*?)_", r"<i>\1</i>", html)

    # 列表
    html = re.sub(r"^\- (.*)$", r"<li>\1</li>", html, flags=re.MULTILINE)
    html = re.sub(r"^\d+\. (.*)$", r"<li>\1</li>", html, flags=re.MULTILINE)

    # 换行
    html = html.replace("\n", "<br/>")

    return html
